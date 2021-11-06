import docx
from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Mm
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import date


def insertHR(paragraph, size, space):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
                              'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                              'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
                              'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
                              'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                              'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
                              'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
                              'w:pPrChange'
                              )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), str(space))
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)


def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existance, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def get_date():
    data = str("Milano, " + date.today().strftime("%d %B %Y"))

    if "January" in data:
        data = data.replace("January", "Gennaio")
    elif "February" in data:
        data = data.replace("February", "Febbraio")
    elif "March" in data:
        data = data.replace("March", "Marzo")
    elif "April" in data:
        data = data.replace("April", "Aprile")
    elif "May" in data:
        data = data.replace("May", "Maggio")
    elif "June" in data:
        data = data.replace("June", "Giugno")
    elif "July" in data:
        data = data.replace("July", "Luglio")
    elif "August" in data:
        data = data.replace("August", "Agosto")
    elif "September" in data:
        data = data.replace("September", "Settembre")
    elif "October" in data:
        data = data.replace("October", "Ottobre")
    elif "November" in data:
        data = data.replace("November", "Novembre")
    elif "December" in data:
        data = data.replace("December", "Dicembre")

    return data


def create_receipt(item_list, numero, causale):
    data = get_date()

    d = Document()

    section = d.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)

    style = d.styles['Normal']
    font = style.font
    font.name = 'Garamond'
    font.size = Pt(10)

    p1 = d.add_paragraph()
    p1.paragraph_format.line_spacing = 1
    p1.alignment = 0
    p1.style = d.styles['Normal']
    p1.add_run("Sara Ezelina Vantini\n").bold = True
    p1.add_run("via Don Carlo Gnocchi 19\n")
    p1.add_run("20148 MILANO\n")
    p1.add_run("Codice Fiscale:\n").bold = True
    p1.add_run("VNTSZL74E71F205T\n").bold = True

    p2 = d.add_paragraph()
    p2.paragraph_format.line_spacing = 1
    p2.alignment = 2
    p2.style = d.styles['Normal']
    p2.add_run("Spett.le\t\t\n")
    p2.add_run("la LUSac srl\t\t\n").bold = True
    p2.add_run("via Defendente Sacchi 13\t\n")
    p2.add_run("27100 PAVIA\t\t\n\n")
    p2.add_run("Partita IVA 02778180188\t\n").bold = True

    p3 = d.add_paragraph()
    p3.paragraph_format.line_spacing = 1
    p3.alignment = 0
    p3.style = d.styles['Normal']
    p3.add_run("Ricevuta nr. " + str(numero) + " / " + date.today().strftime("%Y"))
    p3.add_run("\n\n" + data)
    p3.add_run("\n\nRicevo dalla Società la LUSac srl le somme sotto specificate a fronte delle prestazioni rientranti in rapporto di prestazione di lavoro autonomo occasionale*\n")

    table = d.add_table(1, 0, style="Light List")
    table.add_column(Mm(129.2))
    table.add_column(Mm(30))
    table.autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell = table.rows[0].cells
    cell[0].text = "SPECIFICHE"
    cell[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
    cell[1].text = "costo netto cadauno"
    cell[1].horizontal_alignment = WD_TABLE_ALIGNMENT.CENTER
    cell[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER

    row0 = table.rows[0]
    row0.height = Mm(10)

    d.add_paragraph()

    table1 = d.add_table(1, 0)
    table1.add_column(Mm(129.2))
    table1.add_column(Mm(30))
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER

    ###################################################################################

    for i in range(len(item_list) - 1):
        table1.add_row()

    row1 = table1.rows

    totale = 0.0
    for i in range(len(item_list)):
        totale = totale + (float(item_list[i][0]) * float(item_list[i][1]))
        row1[i].cells[0].text = str(item_list[i][0]) + "x " + item_list[i][2]
        row1[i].cells[1].text = "€   %.2f" % (item_list[i][1])
        i = i + 1

    table1.add_row()
    row1 = table1.rows

    p4 = row1[-1].cells[1].add_paragraph()
    p4.alignment = 0
    p4.paragraph_format.line_spacing = 1
    p4.add_run("\n\nTotale €  %.2f" % totale)

    ######################################################################################

    d.add_page_break()

    p5 = d.add_paragraph()
    p5.paragraph_format.line_spacing = 1
    p5.alignment = 0
    p5.style = d.styles['Normal']
    p5.add_run("Sara Ezelina Vantini\n").bold = True
    p5.add_run("via Don Carlo Gnocchi 19\n")
    p5.add_run("20148 MILANO\n")
    p5.add_run("Codice Fiscale:\n").bold = True
    p5.add_run("VNTSZL74E71F205T\n").bold = True

    p6 = d.add_paragraph()
    p6.paragraph_format.line_spacing = 1
    p6.alignment = 2
    p6.style = d.styles['Normal']
    p6.add_run("Spett.le\t\t\t\n")
    p6.add_run("la LUSac srl\t\t\n").bold = True
    p6.add_run("via Defendente Sacchi 13\t\n")
    p6.add_run("27100 PAVIA\t\t\n\n")
    p6.add_run("Partita IVA 02778180188\t\n").bold = True

    p7 = d.add_paragraph()
    p7.paragraph_format.line_spacing = 1
    p7.alignment = 0
    p7.style = d.styles['Normal']
    p7.add_run("Ricevuta nr. " + str(numero) + " / " + date.today().strftime("%Y"))
    p7.add_run("\n\n" + data)
    p7.add_run(
        "\n\nRicevo dalla Società la LUSac srl le somme sotto specificate a fronte delle prestazioni rientranti in rapporto di prestazione di lavoro autonomo occasionale*")

    table2 = d.add_table(4, 0, style="Light List")
    table2.add_column(Mm(129.2))
    table2.add_column(Mm(30))
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    row2 = table2.rows
    row2[0].cells[0].text = "Tipologia della prestazione"
    row2[0].cells[0].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
    row2[0].cells[0].horizontal_alignment = WD_TABLE_ALIGNMENT.CENTER

    row2[0].cells[1].text = "Compenso"
    row2[0].cells[1].vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
    row2[0].cells[1].horizontal_alignment = WD_TABLE_ALIGNMENT.CENTER

    table2.rows[0].height = Mm(10)

    p8 = row2[1].cells[0].add_paragraph()
    p8.paragraph_format.line_spacing = 1
    p8.alignment = 0
    p8.style = d.styles['Normal']
    p8.add_run("\n\n\n")
    p8.add_run(causale)
    p8.add_run("\n\n\n")

    p8 = row2[1].cells[1].add_paragraph()
    p8.paragraph_format.line_spacing = 1
    p8.alignment = 2
    p8.style = d.styles['Normal']
    p8.add_run("\n\n\n")
    p8.add_run("€   %.2f" % (totale * 1.25))
    p8.add_run("\n\n\n")

    p8 = row2[2].cells[0].add_paragraph()
    p8.paragraph_format.line_spacing = 1
    p8.alignment = 2
    p8.style = d.styles['Normal']
    p8.add_run("Ritenuta d'acconto 20%")
    p8.add_run("\n")

    p8 = row2[2].cells[1].add_paragraph()
    p8.paragraph_format.line_spacing = 1
    p8.alignment = 2
    p8.style = d.styles['Normal']
    p8.add_run("€   %.2f" % (totale * 0.25))
    p8.add_run("\n")

    p8 = row2[3].cells[0].add_paragraph()
    p8.paragraph_format.line_spacing = 1
    p8.alignment = 2
    p8.style = d.styles['Normal']
    p8.add_run("Netto a pagare")
    p8.add_run("\n")

    p8 = row2[3].cells[1].add_paragraph()
    p8.paragraph_format.line_spacing = 1
    p8.alignment = 2
    p8.style = d.styles['Normal']
    p8.add_run("€   %.2f" % totale)
    p8.add_run("\n")

    p9 = d.add_paragraph()
    p9.paragraph_format.line_spacing = 1
    p9.alignment = 0
    p9.style = d.styles['Normal']
    p9.add_run("\nTale importo dovrà essere bonificato sul seguente conto corrente\n")
    p9.add_run("Banco BPM\n")
    p9.add_run("Iban: IT07H0503401753000000101618\n")
    p9.add_run("Swift: BPMIITMMXXX\n")
    p9.add_run("intestatario: Sara Ezelina Vantini\n")

    # adding space between table and paragraphs
    d.add_paragraph().add_run("\n\n\n\n\n")

    p10 = d.add_paragraph()
    p10.paragraph_format.line_spacing = 1
    p10.alignment = 0
    p10.style = d.styles['Normal']
    p10.add_run("*Le eventuali marche da bollo per la contabilità sono a carico della stessa ditta committente.\n")
    p10.add_run("La prestazione oggetto della presente nota è stata effettuata in via occasionale, contingente ed episodica; il relativo compenso è da inquadrare tra i redditi di cui all’art. 67 comma 1, lettera L, del D.P.R. 917/86 e, pertanto, esclusa dal campo di applicazione dell’I.V.A. ai sensi dell'art. 5 del D.P.R. n. 633 del 26 ottobre 1972 e successive modifiche ed integrazioni.")

    return d

# item_list = [[3, 5.0, 'Shopper upcycling'], [2, 4.0, 'Coppia lettere UNICA'], [10, 2.0, 'Sportina con tasca'],[6, 7.0, 'test1'],[9, 5.6, 'test2'],[1, 9.9, 'test3']]
# causale = "Confezione di n. 5 borse LUSac"
# numero = 5
# create_receipt(item_list, numero, causale)
